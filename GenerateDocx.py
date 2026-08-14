from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt



def constructFile(student_name, start_date, 
                  mandatories, numMandatories, mandatoryCourses,
                  seminar, 
                  internalElectives, externalElectives, numElectives):
    doc = Document()
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.35

    # titles
    title = doc.add_paragraph()
    addRun(title,'בדיקת השלמת חובות שמיעה בחטיבה לניהול וישוב סכסוכים',True,True,alignment= WD_PARAGRAPH_ALIGNMENT.CENTER,size=16)

    smallTitle = doc.add_paragraph()
    addRun(smallTitle,'סה"כ יש להשלים 28 נק"ז',True,alignment= WD_PARAGRAPH_ALIGNMENT.CENTER)

    # details
    nameP = doc.add_paragraph()
    addRun(nameP,'שם הסטודנט/ית: ',True)
    addRun(nameP,f'       {student_name}       ',underline=True)

    dateP = doc.add_paragraph()
    addRun(dateP,'תחילת לימודים בחטיבה: ',True)
    addRun(dateP,f'       {start_date}       ',underline=True)

    addSpacer(doc)

    # mandatory courses
    mandatoryP = doc.add_paragraph()
    addRun(mandatoryP,'קורסי החובה בחטיבה 8 נק"ז - סה"ז נק"ז חובה שהושלמו: ',True,True)
    addRun(mandatoryP,f'       {numMandatories}       ',underline=True)

    addMandatories(doc,mandatories,mandatoryCourses)

    addSpacer(doc)

    # seminar

    seminarP = doc.add_paragraph()
    addRun(seminarP,'סמינר בשנה ג 4 נק"ז (אחד מבין השניים) - כל אחד קורס שנתי בהיקף של 4 נק"ז',True,True)

    addSeminar(doc,seminar)

    addSpacer(doc)

    # electives

    electivesP = doc.add_paragraph()
    addRun(electivesP,'קורסי הבחירה בחטיבה 16 נק"ז - סה"כ נק"ז בחירה שהושלמו: ',True,True)
    addRun(electivesP,f'       {numElectives}       ',underline=True)

    #addSpacer(doc)

    addElectivesTable(doc,internalElectives,externalElectives)

    addSpacer(doc)

    # summary
    summaryFont = "Arial"
    summaryTitle = doc.add_paragraph()
    addRun(summaryTitle,'סיכום:',True,True,alignment= WD_PARAGRAPH_ALIGNMENT.CENTER,font=summaryFont)

    overall = doc.add_paragraph()
    if seminar!=0:
        semCred = 4
    else:
        semCred = 0
    sum = numMandatories+numElectives+semCred
    addRun(overall,'סה"כ: חובה ',font=summaryFont)
    addRun(overall,f'  {numMandatories}  ',underline=True,font=summaryFont)
    addRun(overall,' + בחירה ',font=summaryFont)
    addRun(overall,f'  {numElectives}  ',underline=True,font=summaryFont)
    addRun(overall,' + סמינר ',font=summaryFont)
    addRun(overall,f'  {semCred}  ',underline=True,font=summaryFont)
    addRun(overall,' = ',font=summaryFont)
    addRun(overall,f'  {sum}  ',underline=True,font=summaryFont)
    addRun(overall,' נק"ז הושלמו בחטיבה מתוך 28 נק"ז נדרשים.',font=summaryFont)

    reqTitle = doc.add_paragraph()
    addRun(reqTitle,'לצורך סגירת התואר יש להשלים:',True,True,font=summaryFont)

    counter = 0

    if numMandatories<8:
        summary1 = doc.add_paragraph()
        counter+=1
        addRun(summary1,f'{counter}.  קורסי חובה: מספר נק"ז ',font=summaryFont)
        addRun(summary1,f'  {8-numMandatories}  ', False,True,font=summaryFont)
        addRun(summary1,'שמות הקורסים: ',font=summaryFont)
        text = ''
        for i in range(4):
            if i not in mandatories:
                text += f'{mandatoryCourses[i]}, '
        text = text[:-2]
        addRun(summary1,text,underline=True,font=summaryFont)

    
    if numElectives<16:
        counter+=1
        summary2 = doc.add_paragraph()
        addRun(summary2,f'{counter}.  קורסי בחירה: מספר נק"ז ',font=summaryFont)
        addRun(summary2,f'  {16-numElectives}  ', False,True,font=summaryFont)

    if seminar == 0:
        counter+=1
        summary3 = doc.add_paragraph()
        addRun(summary3,f'{counter}.  סמינר - 4 נק"ז',font=summaryFont)


    doc.save(f'{student_name} סיכום חובות קורסים.docx')

def addElectivesTable(doc,internalElectives, externalElectives):
    rowNum = max(len(internalElectives),len(externalElectives))
    table = doc.add_table(rows=rowNum+1, cols=4)
    set_table_borders(table)
    set_table_rtl(table)

    widths = [Inches(2.3), Inches(0.5), Inches(2.7), Inches(0.5)]
    set_col_widths(table, widths)

    table.cell(0,0).text = 'קורסי בחירה מבין קורסי החטיבה'
    table.cell(0,1).text = 'נק"ז'
    table.cell(0,2).text = 'קורסי בחירה מאושרים ממחלקות אחרות'
    table.cell(0,3).text = 'נק"ז'

    for row,(name,credits) in zip(table.rows[1:],internalElectives):
        row.cells[0].text = name
        row.cells[1].text = credits

    for row,(name,credits) in zip(table.rows[1:],externalElectives):
        row.cells[2].text = name
        row.cells[3].text = credits

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in paragraph.runs:
                make_run_rtl(run,True)

    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                for run in paragraph.runs:
                    make_run_rtl(run)

def set_col_widths(table, widths):
    # widths: list of Inches() values, one per column
    table.autofit = False

    # Force fixed layout
    tblPr = table._tbl.tblPr
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Set the grid column widths
    for i, width in enumerate(widths):
        table.columns[i].width = width

    # Set width on every individual cell (this is the part that's usually missing)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

def set_table_rtl(table):
    tblPr = table._tbl.tblPr

    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")

    tblPr.append(bidi)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)

    tblPr.append(borders)

def addSeminar(doc,seminar):
    items = [
        'התמודדות מתבגרים במצבי קונפליט',
        'קונפליקטים במרחב המשפחתי'
    ]
    if seminar == 0:
        sem1 = doc.add_paragraph()
        sem1.paragraph_format.right_indent = Inches(0.5)
        addRun(sem1,"☐ ")
        addRun(sem1,items[0])

        sem2 = doc.add_paragraph()
        sem2.paragraph_format.right_indent = Inches(0.5)
        addRun(sem2,"☐ ")
        addRun(sem2,items[1])
    else:
        for i in range(2):
            sem = doc.add_paragraph()
            sem.paragraph_format.right_indent = Inches(0.5)
            if i+1 == seminar:
                addRun(sem,"☑ ")
                addRun(sem,items[i])
            else:
                addRun(sem,"☐ ")
                addRun(sem,items[i])


def addMandatories(doc,mandatories,mandatoryCourses):
    for i in range(4):
        mand = doc.add_paragraph()
        mand.paragraph_format.right_indent = Inches(0.5)
        if i in mandatories:
            addRun(mand,"☑ ")
            addRun(mand,mandatoryCourses[i]+ ' - 2 נק"ז')
        else:
            addRun(mand,"☐ ")
            addRun(mand,mandatoryCourses[i] + ' - 2 נק"ז')

def addSpacer(doc):
    empty = doc.add_paragraph()
    addRun(empty,"")

def addRun(paragraph, text, bold=False, underline=False,alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT,size=12,font="David"):
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    make_run_rtl(run,bold,underline,size,font)

def make_run_rtl(run, bold=False, underline=False,size=12,font="David"):
    r = run._element
    rPr = r.get_or_add_rPr()

    # RTL
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

    # Bold
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    # Font size
    if size is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))
        rPr.append(sz)

        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)

    # Font
    if font is not None:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rFonts.set(qn('w:cs'), font)
        rPr.append(rFonts)

#constructFile("גיא אברהם","23.8.2024",[1,3],4,[
#        'מושגי יסוד בחקר סכסוכים',
#        'מבנה חברתי של ישראל',
#        'סוגיות בקונפליקטים קבוצתיים וארגוניים',
#        'משא ומתן ככלי לניהול וישוב סכסוכים'
#    ],2,[("בחירה 1","3"),("בחירה 2", "5")],[("בחירה 1","3"),("בחירה 1","3"),("בחירה 1","3")],10)
